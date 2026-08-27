import ssl
import eventlet
eventlet.monkey_patch()

import os
import uuid
import json
import zipfile
import io
from celery import Celery
from pypdf import PdfReader
import docx
from dotenv import load_dotenv # <-- NEW IMPORT

# Load the secrets from the .env file
load_dotenv(r"C:/pii-redaction-webapp/backend/venv/app/.env")

# Import our NLP engine
from app.services.presidio_service import redaction_service

# Fetch the URL securely from the environment
REDIS_URL = os.getenv("REDIS_URL")

celery_app = Celery("redaction_worker", broker=REDIS_URL, backend=REDIS_URL)
celery_app.conf.broker_use_ssl = {'ssl_cert_reqs': ssl.CERT_NONE}
celery_app.conf.redis_backend_use_ssl = {'ssl_cert_reqs': ssl.CERT_NONE}

# --- BULLETPROOF PATH RESOLUTION ---
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = r"C:/pii-redaction-webapp/backend/temp_outputs"
os.makedirs(TEMP_DIR, exist_ok=True)

@celery_app.task
def cleanup_file_task(file_path: str):
    """Deletes a file after its TTL expires."""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"[CLEANUP] Deleted expired file: {file_path}")
    except Exception as e:
        print(f"[CLEANUP ERROR] Could not delete {file_path}: {e}")

@celery_app.task
def process_batch_task(file_metadata, entity_list, mask_char):
    """Reads saved files from disk, runs NLP redaction, and builds a ZIP."""
    batch_results = []
    generated_file_paths = []
    
    for file_data in file_metadata:
        file_path = file_data["path"]
        filename = file_data["filename"]
        extension = filename.split('.')[-1].lower()
        extracted_text = ""
        
        try:
            # 1. Read from Disk
            with open(file_path, "rb") as f:
                content = f.read()
                
            if extension == "txt":
                extracted_text = content.decode("utf-8")
            elif extension == "json":
                extracted_text = json.dumps(json.loads(content), indent=2) 
            elif extension == "pdf":
                pdf_reader = PdfReader(io.BytesIO(content))
                for page in pdf_reader.pages:
                    extracted_text += page.extract_text() + "\n"
            elif extension == "docx":
                doc = docx.Document(io.BytesIO(content))
                extracted_text = "\n".join([para.text for para in doc.paragraphs])

            # 2. Process Text
            result = redaction_service.process_text(
                text=extracted_text, language="en", entities=entity_list, mask_char=mask_char
            )
            
            # 3. Rebuild File
            safe_id = uuid.uuid4().hex[:8]
            new_filename = f"redacted_{safe_id}_{filename}"
            if extension == "pdf":
                new_filename = new_filename.replace('.pdf', '.txt')
                
            output_path = os.path.join(TEMP_DIR, new_filename) # <-- Use global TEMP_DIR
            
            if extension in ["txt", "json", "pdf"]:
                with open(output_path, "w", encoding="utf-8") as out_f:
                    out_f.write(result["redacted_text"])
            elif extension == "docx":
                new_doc = docx.Document()
                for line in result["redacted_text"].split('\n'):
                    if line.strip():
                        new_doc.add_paragraph(line)
                new_doc.save(output_path)
                
            generated_file_paths.append(output_path)

            # Schedule individual file cleanup 30 minutes (1800s) in the future
            cleanup_file_task.apply_async(args=[output_path], countdown=1800)
            
            # 4. Clean up the raw input file from the disk
            if os.path.exists(file_path):
                os.remove(file_path)

            batch_results.append({
                "filename": filename,
                "download_url": f"http://localhost:8000/api/v1/download/{new_filename}",
                "total_entities_found": result["total_entities_found"],
                "entities_detected": result["entities_detected"]
            })
            
        except Exception as e:
            print(f"Worker failed to process {filename}: {e}")

            
# 5. Create ZIP
    zip_download_url = None
    if generated_file_paths:
        zip_filename = f"batch_redacted_{uuid.uuid4().hex[:8]}.zip"
        
        zip_path = os.path.join(TEMP_DIR, zip_filename) # <-- Use global TEMP_DIR
        
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for path in generated_file_paths:
                zipf.write(path, os.path.basename(path))
        zip_download_url = f"http://localhost:8000/api/v1/download/{zip_filename}"

        # Schedule ZIP cleanup 30 minutes in the future
        cleanup_file_task.apply_async(args=[zip_path], countdown=1800)
            
    return {
        "batch_results": batch_results,
        "zip_download_url": zip_download_url
    }