import os
import uuid
from fastapi.responses import FileResponse
from fastapi import FastAPI, HTTPException, File, UploadFile, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from typing import Optional, List
import io
import json
from pypdf import PdfReader
import docx
import zipfile
import time

from celery.result import AsyncResult
from app.worker import process_batch_task, celery_app

# Import our custom schemas and the Presidio service
from app.models.schemas import RedactionRequest, RedactionResponse
from app.services.presidio_service import redaction_service

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMP_DIR = r"C:/pii-redaction-webapp/backend/temp_outputs"
os.makedirs(TEMP_DIR, exist_ok=True)

def cleanup_file(path: str):
    """Deletes a file from the disk to prevent storage leaks."""
    try:
        if os.path.exists(path):
            os.remove(path)
            print(f"Cleaned up temporary file: {path}")
    except Exception as e:
        print(f"Error cleaning up file {path}: {e}")

app = FastAPI(
    title="PII Redaction API",
    description="Enterprise API for text anonymization and PII detection",
    version="1.0.0"
)
# Create a directory to hold the generated files
os.makedirs("temp_outputs", exist_ok=True)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"], 
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

TEMP_DIR = os.path.abspath("temp_outputs")

@app.get("/api/v1/download/{filename}")
def download_file(filename: str):
    file_path = os.path.join(TEMP_DIR, filename)

    # 1. Verify existence
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found or expired.")

    # 2. Check 30-minute TTL (1800 seconds)
    file_age_seconds = time.time() - os.path.getmtime(file_path)
    if file_age_seconds > 1800:
        try:
            os.remove(file_path)
        except OSError:
            pass
        raise HTTPException(status_code=410, detail="Download link has expired (30-minute limit reached).")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/octet-stream"
    )

@app.get("/")
def health_check():
    return {
        "status": 200, 
        "message": "PII Redaction API is running."
    }

@app.post("/api/v1/redact", response_model=RedactionResponse)
def redact_text(request: RedactionRequest):
    """
    Analyzes the input text, detects PII entities, and returns the redacted string along with metadata.
    """
    try:
        # Pass the validated data from the Pydantic schema directly into our service
        result = redaction_service.process_text(
            text=request.text,
            language=request.language,
            entities=request.entities,
            mask_char=request.mask_char
        )
        return result
    except Exception as e:
        # If anything goes wrong inside the Presidio engine, catch it and return a clean 500 error
        raise HTTPException(status_code=500, detail=f"An error occurred during processing: {str(e)}")

# --- NEW FILE UPLOAD ENDPOINT ---

# @app.post("/api/v1/redact/files")
# async def redact_files(
#     files: List[UploadFile] = File(...),
#     entities: Optional[str] = Form(None),
#     mask_char: str = Form("*")
# ):
#     """Extracts and redacts PII from multiple uploaded files, returning a batch summary and a ZIP archive."""
#     batch_results = []
#     generated_file_paths = [] # Track paths to zip them later
    
#     entity_list = None
#     if entities:
#         entity_list = [e.strip() for e in entities.split(",") if e.strip()]

#     for file in files:
#         extension = file.filename.split('.')[-1].lower()
#         extracted_text = ""
        
#         try:
#             # 1. Extract Text
#             if extension == "txt":
#                 content = await file.read()
#                 extracted_text = content.decode("utf-8")
#             elif extension == "json":
#                 content = await file.read()
#                 extracted_text = json.dumps(json.loads(content), indent=2) 
#             elif extension == "pdf":
#                 content = await file.read()
#                 pdf_reader = PdfReader(io.BytesIO(content))
#                 for page in pdf_reader.pages:
#                     extracted_text += page.extract_text() + "\n"
#             elif extension == "docx":
#                 content = await file.read()
#                 doc = docx.Document(io.BytesIO(content))
#                 extracted_text = "\n".join([para.text for para in doc.paragraphs])
#             else:
#                 continue 

#             # 2. Process Text
#             result = redaction_service.process_text(
#                 text=extracted_text, language="en", entities=entity_list, mask_char=mask_char
#             )
            
#             # 3. Rebuild File
#             safe_id = uuid.uuid4().hex[:8]
#             new_filename = f"redacted_{safe_id}_{file.filename}"
#             if extension == "pdf":
#                 new_filename = new_filename.replace('.pdf', '.txt')
                
#             output_path = os.path.join("temp_outputs", new_filename)
            
#             if extension in ["txt", "json", "pdf"]:
#                 with open(output_path, "w", encoding="utf-8") as f:
#                     f.write(result["redacted_text"])
#             elif extension == "docx":
#                 new_doc = docx.Document()
#                 for line in result["redacted_text"].split('\n'):
#                     if line.strip():
#                         new_doc.add_paragraph(line)
#                 new_doc.save(output_path)
                
#             generated_file_paths.append(output_path)

#             # 4. Append specific file metrics to the batch
#             batch_results.append({
#                 "filename": file.filename,
#                 "download_url": f"http://localhost:8000/api/v1/download/{new_filename}",
#                 "total_entities_found": result["total_entities_found"],
#                 "entities_detected": result["entities_detected"]
#             })
            
#         except Exception as e:
#             print(f"Failed to process {file.filename}: {e}")
            
#     # 5. Create a ZIP file if anything was generated
#     zip_download_url = None
#     if generated_file_paths:
#         zip_filename = f"batch_redacted_{uuid.uuid4().hex[:8]}.zip"
#         zip_path = os.path.join("temp_outputs", zip_filename)
        
#         with zipfile.ZipFile(zip_path, 'w') as zipf:
#             for path in generated_file_paths:
#                 zipf.write(path, os.path.basename(path))
                
#         zip_download_url = f"http://localhost:8000/api/v1/download/{zip_filename}"
            
#     return {
#         "batch_results": batch_results,
#         "zip_download_url": zip_download_url
#     }


@app.post("/api/v1/redact/files")
async def redact_files_async(
    files: List[UploadFile] = File(...),
    entities: Optional[str] = Form(None),
    mask_char: str = Form("*")
):
    """Saves raw files to disk and drops a redaction task into the Celery/Redis queue."""
    file_metadata = []
    
    for file in files:
        safe_id = uuid.uuid4().hex[:8]
        temp_path = os.path.join(TEMP_DIR, f"raw_{safe_id}_{file.filename}") # <-- Updated to TEMP_DIR
        
        with open(temp_path, "wb") as buffer:
            buffer.write(await file.read())
            
        file_metadata.append({"path": temp_path, "filename": file.filename})

    # 2. Parse entities
    entity_list = None
    if entities:
        entity_list = [e.strip() for e in entities.split(",") if e.strip()]

    # 3. Dispatch the Celery task
    task = process_batch_task.delay(file_metadata, entity_list, mask_char)
    
    return {"task_id": str(task.id), "message": "Batch processing started."}

@app.get("/api/v1/tasks/{task_id}")
def get_task_status(task_id: str):
    """Endpoint for React to poll and check if Celery is finished processing."""
    task_result = AsyncResult(task_id, app=celery_app)
    
    if task_result.state == 'SUCCESS':
        return {"status": task_result.state, "result": task_result.result}
    elif task_result.state == 'FAILURE':
        return {"status": task_result.state, "error": str(task_result.info)}
    
    # E.g., 'PENDING', 'STARTED', or 'RETRY'
    return {"status": task_result.state}